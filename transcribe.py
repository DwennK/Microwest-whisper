from __future__ import annotations

import argparse
import json
import os
import platform
import re
import shutil
import subprocess
import sys
import unicodedata
import warnings
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from license_client import activate_license as activate_product_license
from license_client import license_status_text, validate_license as validate_product_license
from transcript_paths import SUPPORTED_AUDIO_EXTENSIONS, transcript_stem, work_wav_path


try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

if load_dotenv:
    load_dotenv()


@dataclass(frozen=True)
class Paths:
    audio: Path
    output_dir: Path
    work_dir: Path
    wav: Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="High-quality French transcription with WhisperX alignment and speaker diarization."
    )
    parser.add_argument("--audio", help="Path to the source audio file, for example recording.m4a.")
    parser.add_argument("--doctor", action="store_true", help="Check Python, FFmpeg, optional GPU and installed packages.")
    parser.add_argument("--check-token", action="store_true", help="Check whether the Hugging Face token can load pyannote.")
    parser.add_argument("--activate-license", default="", help="Activate a Microwest Whisper license key on this machine.")
    parser.add_argument("--license-status", action="store_true", help="Show the current Microwest Whisper license status.")
    parser.add_argument("--output-dir", default="output", help="Directory where transcript files are written.")
    parser.add_argument("--work-dir", default="work", help="Directory for intermediate WAV files.")
    parser.add_argument(
        "--profile",
        choices=["manual", "auto", "quality", "fast", "cpu", "no-speakers"],
        default="manual",
        help="Preset profile. manual keeps the explicit options.",
    )
    parser.add_argument("--model", default="large-v3", help="Whisper model. Use large-v3 for best quality.")
    parser.add_argument(
        "--asr-backend",
        choices=["auto", "whisperx", "mlx"],
        default="auto",
        help="ASR engine. auto uses MLX on Apple Silicon when installed, otherwise WhisperX/faster-whisper.",
    )
    parser.add_argument(
        "--mlx-model",
        default="auto",
        help="MLX model repo/path. auto maps large-v3 to mlx-community/whisper-large-v3-mlx.",
    )
    parser.add_argument("--language", default="fr", help="Input language code. French is 'fr'.")
    parser.add_argument("--batch-size", type=int, default=8, help="WhisperX batch size. Lower this if RAM is tight.")
    parser.add_argument(
        "--threads",
        type=int,
        default=0,
        help="CPU threads for faster-whisper. 0 chooses a sensible automatic value.",
    )
    parser.add_argument("--device", choices=["auto", "cpu", "cuda"], default="auto")
    parser.add_argument("--compute-type", default="auto", help="auto, float16, int8, int8_float16, float32.")
    parser.add_argument("--hf-token", default=os.environ.get("HUGGINGFACE_TOKEN"), help="Hugging Face token for pyannote.")
    parser.add_argument("--no-diarization", action="store_true", help="Transcribe without speaker diarization.")
    parser.add_argument("--speakers", type=int, default=None, help="Exact number of speakers, if known.")
    parser.add_argument("--min-speakers", type=int, default=None, help="Minimum number of speakers.")
    parser.add_argument("--max-speakers", type=int, default=None, help="Maximum number of speakers.")
    parser.add_argument("--delete-work", action="store_true", help="Delete the intermediate WAV file after success.")
    parser.add_argument("--force", action="store_true", help="Ignore reusable checkpoints and recompute all stages.")
    parser.add_argument(
        "--audio-filter",
        choices=["loudnorm", "voice-clean", "none"],
        default="loudnorm",
        help="FFmpeg audio filter before ASR.",
    )
    parser.add_argument("--trim-silence", action="store_true", help="Trim leading and trailing silence during preprocessing.")
    parser.add_argument(
        "--speaker-map",
        default="",
        help='Rename speakers, for example "SPEAKER_00=Alice,SPEAKER_01=Bob".',
    )
    parser.add_argument("--speaker-map-file", default="", help="JSON file containing speaker rename mappings.")
    parser.add_argument(
        "--rename-only",
        action="store_true",
        help="Regenerate exports from the saved final checkpoint after applying speaker names.",
    )
    return parser.parse_args()


def fail(message: str) -> None:
    print(f"ERROR: {message}", file=sys.stderr)
    raise SystemExit(1)


def apply_profile(args: argparse.Namespace) -> None:
    profile = args.profile
    if profile == "manual":
        return
    if profile == "auto":
        profile = recommended_profile()

    if profile == "fast":
        args.model = "large-v3-turbo"
        args.batch_size = max(args.batch_size, 12)
        args.device = "auto"
        args.compute_type = "auto"
    elif profile == "cpu":
        args.model = "medium"
        args.batch_size = min(args.batch_size, 4)
        args.device = "cpu"
        args.compute_type = "int8"
    elif profile == "no-speakers":
        args.model = "large-v3"
        args.no_diarization = True
    elif profile == "quality":
        args.model = "large-v3"
        args.batch_size = 8
        args.device = "auto"
        args.compute_type = "auto"


def recommended_profile() -> str:
    if is_apple_silicon():
        return "quality"
    try:
        import torch

        if torch.cuda.is_available():
            return "quality"
    except Exception:
        pass
    cpu_count = os.cpu_count() or 4
    return "cpu" if cpu_count <= 8 else "fast"


def validate_args(args: argparse.Namespace) -> None:
    if args.doctor:
        return
    if getattr(args, "license_status", False) or getattr(args, "activate_license", ""):
        return
    if args.check_token and not args.audio:
        return
    if not args.audio:
        fail("--audio is required unless --doctor or --check-token is used alone.")

    speaker_values = {
        "--speakers": args.speakers,
        "--min-speakers": args.min_speakers,
        "--max-speakers": args.max_speakers,
    }
    for option, value in speaker_values.items():
        if value is not None and value < 1:
            fail(f"{option} must be greater than 0.")

    if args.speakers and (args.min_speakers or args.max_speakers):
        fail("Use either --speakers or --min-speakers/--max-speakers, not both.")
    if args.min_speakers and args.max_speakers and args.min_speakers > args.max_speakers:
        fail("--min-speakers must be lower than or equal to --max-speakers.")
    if args.rename_only and not args.speaker_map and not args.speaker_map_file:
        fail("--rename-only requires --speaker-map or --speaker-map-file.")


def build_paths(args: argparse.Namespace) -> Paths:
    audio = Path(args.audio).expanduser().resolve()
    if not audio.exists():
        fail(f"Audio file not found: {audio}")
    if audio.suffix.lower() not in SUPPORTED_AUDIO_EXTENSIONS:
        fail(
            f"Unsupported audio extension '{audio.suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_AUDIO_EXTENSIONS))}"
        )

    output_dir = Path(args.output_dir).resolve()
    work_dir = Path(args.work_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    work_dir.mkdir(parents=True, exist_ok=True)
    wav = work_wav_path(audio, work_dir)
    return Paths(audio=audio, output_dir=output_dir, work_dir=work_dir, wav=wav)


def checkpoint_path(paths: Paths, stage: str) -> Path:
    return paths.work_dir / f"{transcript_stem(paths.audio)}.{stage}.json"


def wav_metadata_path(paths: Paths) -> Path:
    return paths.wav.with_suffix(paths.wav.suffix + ".meta.json")


def audio_preprocess_signature(args: argparse.Namespace) -> dict[str, Any]:
    return {
        "version": 1,
        "audio_filter": args.audio_filter,
        "trim_silence": bool(args.trim_silence),
    }


def checkpoint_signature(args: argparse.Namespace, stage: str, resolved_backend: str) -> dict[str, Any]:
    signature: dict[str, Any] = {
        "version": 1,
        "stage": stage,
        "audio": audio_preprocess_signature(args),
        "profile": args.profile,
        "model": args.model,
        "language": args.language,
        "asr_backend": args.asr_backend,
        "resolved_asr_backend": resolved_backend,
        "mlx_model": resolve_mlx_model(args.model, args.mlx_model) if resolved_backend == "mlx" else None,
        "batch_size": args.batch_size,
        "threads": args.threads,
        "device": args.device,
        "compute_type": args.compute_type,
    }
    if stage.startswith("final."):
        signature.update(
            {
                "no_diarization": bool(args.no_diarization),
                "speakers": args.speakers,
                "min_speakers": args.min_speakers,
                "max_speakers": args.max_speakers,
            }
        )
    return signature


def checkpoint_matches_settings(checkpoint: dict[str, Any], signature: dict[str, Any] | None) -> bool:
    if signature is None:
        return True
    metadata = checkpoint.get("_checkpoint")
    if not isinstance(metadata, dict):
        return False
    return metadata.get("signature") == signature


def load_checkpoint(
    paths: Paths,
    stage: str,
    force: bool = False,
    signature: dict[str, Any] | None = None,
) -> dict[str, Any] | None:
    path = checkpoint_path(paths, stage)
    if force or not path.exists() or path.stat().st_mtime < paths.audio.stat().st_mtime:
        return None
    with path.open("r", encoding="utf-8") as handle:
        checkpoint = json.load(handle)
    if not checkpoint_matches_settings(checkpoint, signature):
        print(f"Ignoring existing {stage} checkpoint because settings changed: {path}")
        return None
    print(f"Using existing {stage} checkpoint: {path}")
    return checkpoint


def save_checkpoint(
    paths: Paths,
    stage: str,
    result: dict[str, Any],
    signature: dict[str, Any] | None = None,
) -> None:
    path = checkpoint_path(paths, stage)
    payload = dict(result)
    if signature is not None:
        payload["_checkpoint"] = {"stage": stage, "signature": signature}
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)
    print(f"Saved {stage} checkpoint: {path}")


def run(command: list[str]) -> None:
    print("+ " + " ".join(f'"{part}"' if " " in part else part for part in command))
    subprocess.run(command, check=True)


def find_ffmpeg() -> str | None:
    ffmpeg = shutil.which("ffmpeg")
    if ffmpeg:
        return ffmpeg

    local_app_data = os.environ.get("LOCALAPPDATA")
    if local_app_data:
        package_root = Path(local_app_data) / "Microsoft" / "WinGet" / "Packages"
        if package_root.exists():
            matches = sorted(package_root.rglob("ffmpeg.exe"))
            if matches:
                ffmpeg_path = str(matches[0])
                os.environ["PATH"] = f"{str(matches[0].parent)}{os.pathsep}{os.environ.get('PATH', '')}"
                return ffmpeg_path
    try:
        import imageio_ffmpeg

        bundled_ffmpeg = Path(imageio_ffmpeg.get_ffmpeg_exe())
        if bundled_ffmpeg.exists():
            os.environ["PATH"] = f"{str(bundled_ffmpeg.parent)}{os.pathsep}{os.environ.get('PATH', '')}"
            return str(bundled_ffmpeg)
    except Exception:
        pass
    return None


def ensure_ffmpeg() -> str:
    ffmpeg = find_ffmpeg()
    if not ffmpeg:
        if platform.system() == "Windows":
            install_hint = "winget install --id Gyan.FFmpeg -e"
        elif platform.system() == "Darwin":
            install_hint = "brew install ffmpeg"
        else:
            install_hint = "install ffmpeg with your package manager"
        fail(f"FFmpeg is not available in PATH. Install it with: {install_hint}")
    return ffmpeg


def build_audio_filter(args: argparse.Namespace) -> str | None:
    filters: list[str] = []
    if args.trim_silence:
        filters.append("silenceremove=start_periods=1:start_duration=0.2:start_threshold=-45dB")
    if args.audio_filter == "loudnorm":
        filters.append("loudnorm=I=-16:TP=-1.5:LRA=11")
    elif args.audio_filter == "voice-clean":
        filters.extend(["highpass=f=80", "lowpass=f=7800", "afftdn=nf=-25", "loudnorm=I=-16:TP=-1.5:LRA=11"])
    return ",".join(filters) if filters else None


def prepare_wav(paths: Paths, args: argparse.Namespace) -> None:
    ffmpeg = ensure_ffmpeg()
    expected_signature = audio_preprocess_signature(args)
    metadata_path = wav_metadata_path(paths)
    if not args.force and paths.wav.exists() and paths.wav.stat().st_mtime >= paths.audio.stat().st_mtime:
        try:
            with metadata_path.open("r", encoding="utf-8") as handle:
                metadata = json.load(handle)
        except (FileNotFoundError, json.JSONDecodeError):
            metadata = {}
        if metadata.get("signature") == expected_signature:
            print(f"Using existing preprocessed WAV: {paths.wav}")
            return
        print(f"Rebuilding preprocessed WAV because audio settings changed: {paths.wav}")

    print("Preparing clean 16 kHz mono WAV for ASR...")
    command = [
        ffmpeg,
        "-y",
        "-i",
        str(paths.audio),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
    ]
    audio_filter = build_audio_filter(args)
    if audio_filter:
        command.extend(["-af", audio_filter])
    command.extend(["-c:a", "pcm_s16le", str(paths.wav)])
    run(command)
    with metadata_path.open("w", encoding="utf-8") as handle:
        json.dump({"signature": expected_signature}, handle, ensure_ascii=False, indent=2)


def select_device(requested: str) -> str:
    if requested != "auto":
        return requested
    try:
        import torch

        return "cuda" if torch.cuda.is_available() else "cpu"
    except Exception:
        return "cpu"


def select_compute_type(device: str, requested: str) -> str:
    if requested != "auto":
        return requested
    return "float16" if device == "cuda" else "int8"


def select_threads(requested: int) -> int:
    if requested > 0:
        return requested
    cpu_count = os.cpu_count() or 4
    return max(4, min(cpu_count - 1, 12))


def is_apple_silicon() -> bool:
    return platform.system() == "Darwin" and platform.machine() == "arm64"


def can_import_mlx_whisper() -> bool:
    try:
        import mlx_whisper  # noqa: F401

        return True
    except Exception:
        return False


def resolve_asr_backend(requested: str) -> str:
    if requested == "mlx":
        if not can_import_mlx_whisper():
            fail("MLX backend requested but mlx-whisper is not installed. On Mac, run ./setup-mac.sh.")
        return "mlx"
    if requested == "whisperx":
        return "whisperx"
    if is_apple_silicon() and can_import_mlx_whisper():
        return "mlx"
    return "whisperx"


def resolve_mlx_model(model: str, requested: str) -> str:
    if requested != "auto":
        return requested
    mapping = {
        "large-v3": "mlx-community/whisper-large-v3-mlx",
        "large-v3-turbo": "mlx-community/whisper-large-v3-turbo",
        "medium": "mlx-community/whisper-medium-mlx",
        "small": "mlx-community/whisper-small-mlx",
        "base": "mlx-community/whisper-base-mlx",
        "tiny": "mlx-community/whisper-tiny",
    }
    return mapping.get(model, f"mlx-community/whisper-{model}-mlx")


def model_language(language: str) -> str | None:
    return None if language == "auto" else language


def is_repetitive_hallucination(text: str) -> bool:
    normalized = strip_accents(text.lower())
    compact = re.sub(r"\s+", "", normalized)
    if len(compact) >= 80:
        for size in range(2, 16):
            repeated = 0
            previous = ""
            for index in range(0, len(compact) - size + 1, size):
                chunk = compact[index : index + size]
                if chunk == previous:
                    repeated += 1
                    if repeated >= 8:
                        return True
                else:
                    repeated = 0
                    previous = chunk

    words = re.findall(r"[\w']+", normalized)
    if len(words) < 8:
        return False

    counts: dict[str, int] = {}
    for word in words:
        counts[word] = counts.get(word, 0) + 1
    most_common = max(counts.values(), default=0)
    if most_common / len(words) >= 0.65:
        return True

    repeated_pairs = 0
    for index in range(len(words) - 2):
        if words[index] == words[index + 1] == words[index + 2]:
            repeated_pairs += 1
    return repeated_pairs >= 3


def sanitize_segments_for_alignment(result: dict[str, Any], source: str) -> dict[str, Any]:
    segments = result.get("segments")
    if not isinstance(segments, list):
        return result

    cleaned = []
    removed = []
    for segment in segments:
        text = " ".join(str(segment.get("text", "")).split()) if isinstance(segment, dict) else ""
        if text and is_repetitive_hallucination(text):
            removed.append(text[:80])
            continue
        cleaned.append(segment)

    if not removed:
        return result

    updated = dict(result)
    updated["segments"] = cleaned
    updated["dropped_repetitive_segments"] = removed
    print(f"Dropped {len(removed)} repetitive ASR segment(s) before {source}.")
    return updated


def transcribe_with_mlx(paths: Paths, args: argparse.Namespace) -> dict[str, Any]:
    import mlx_whisper

    model_name = resolve_mlx_model(args.model, args.mlx_model)
    print(f"ASR backend: mlx-whisper")
    print(f"MLX model: {model_name}")
    print("Transcribing with MLX on Apple Silicon...")
    kwargs: dict[str, Any] = {
        "path_or_hf_repo": model_name,
        "task": "transcribe",
        "word_timestamps": False,
        "verbose": False,
    }
    language = model_language(args.language)
    if language:
        kwargs["language"] = language
    try:
        result = mlx_whisper.transcribe(str(paths.wav), **kwargs)
    except TypeError:
        kwargs.pop("verbose", None)
        result = mlx_whisper.transcribe(str(paths.wav), **kwargs)

    segments = []
    for segment in result.get("segments", []):
        text = " ".join(str(segment.get("text", "")).split())
        if not text:
            continue
        segments.append(
            {
                "start": float(segment.get("start", 0.0)),
                "end": float(segment.get("end", 0.0)),
                "text": text,
            }
        )

    return {
        "segments": segments,
        "language": result.get("language") or model_language(args.language) or "fr",
        "text": result.get("text", ""),
        "asr_backend": "mlx",
        "mlx_model": model_name,
        "source_audio": str(paths.audio),
        "preprocessed_wav": str(paths.wav),
    }


def transcribe_with_whisperx(audio: Any, args: argparse.Namespace) -> tuple[dict[str, Any], str]:
    import gc

    import whisperx

    device = select_device(args.device)
    compute_type = select_compute_type(device, args.compute_type)
    language = model_language(args.language)
    print(f"ASR backend: whisperx/faster-whisper")
    print(f"Device: {device}")
    print(f"Compute type: {compute_type}")
    threads = select_threads(args.threads)
    print(f"CPU threads: {threads}")
    print(f"Model: {args.model}")

    print("Loading ASR model...")
    try:
        model = whisperx.load_model(
            args.model,
            device,
            compute_type=compute_type,
            language=language,
            threads=threads,
        )
    except ValueError as exc:
        if device == "cpu" and compute_type != "float32":
            print(f"Retrying on CPU with float32 after: {exc}")
            compute_type = "float32"
            model = whisperx.load_model(
                args.model,
                device,
                compute_type=compute_type,
                language=language,
                threads=threads,
            )
        else:
            raise

    print("Transcribing...")
    result = model.transcribe(audio, batch_size=args.batch_size, language=language)
    result["asr_backend"] = "whisperx"

    del model
    gc.collect()
    return result, device


def format_ts(seconds: float | None, sep: str = ",") -> str:
    if seconds is None:
        seconds = 0.0
    milliseconds = round(float(seconds) * 1000)
    hours, rest = divmod(milliseconds, 3_600_000)
    minutes, rest = divmod(rest, 60_000)
    secs, millis = divmod(rest, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}{sep}{millis:03d}"


def normalize_segments(result: dict[str, Any]) -> list[dict[str, Any]]:
    segments = result.get("segments", [])
    normalized: list[dict[str, Any]] = []
    for segment in segments:
        text = " ".join(str(segment.get("text", "")).split())
        if not text:
            continue
        normalized.append(
            {
                "start": float(segment.get("start", 0.0)),
                "end": float(segment.get("end", 0.0)),
                "speaker": segment.get("speaker", "SPEAKER_UNKNOWN"),
                "text": text,
                "words": segment.get("words", []),
            }
        )
    return normalized


def merge_turns(segments: list[dict[str, Any]], max_gap: float = 1.2) -> list[dict[str, Any]]:
    turns: list[dict[str, Any]] = []
    for segment in segments:
        if (
            turns
            and turns[-1]["speaker"] == segment["speaker"]
            and segment["start"] - turns[-1]["end"] <= max_gap
        ):
            turns[-1]["end"] = max(turns[-1]["end"], segment["end"])
            turns[-1]["text"] = f'{turns[-1]["text"]} {segment["text"]}'.strip()
        else:
            turns.append({key: segment[key] for key in ("start", "end", "speaker", "text")})
    return turns


def write_txt(path: Path, turns: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for turn in turns:
            handle.write(
                f'[{format_ts(turn["start"], sep=".")} - {format_ts(turn["end"], sep=".")}] '
                f'{turn["speaker"]}: {turn["text"]}\n\n'
            )


def write_srt(path: Path, segments: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for index, segment in enumerate(segments, start=1):
            handle.write(f"{index}\n")
            handle.write(f'{format_ts(segment["start"])} --> {format_ts(segment["end"])}\n')
            handle.write(f'{segment["speaker"]}: {segment["text"]}\n\n')


def write_markdown(path: Path, turns: list[dict[str, Any]], source: Path) -> None:
    with path.open("w", encoding="utf-8") as handle:
        handle.write(f"# Transcription\n\n")
        handle.write(f"Source: `{source.name}`\n\n")
        for turn in turns:
            handle.write(
                f'**{turn["speaker"]}** '
                f'`{format_ts(turn["start"], sep=".")} - {format_ts(turn["end"], sep=".")}`\n\n'
            )
            handle.write(f'{turn["text"]}\n\n')


def write_clean_txt(path: Path, turns: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for turn in turns:
            handle.write(f'{turn["speaker"]}: {turn["text"]}\n\n')


def extract_meeting_notes(turns: list[dict[str, Any]]) -> dict[str, Any]:
    action_pattern = re_compile(
        r"\b(il faut|on doit|je vais|tu vas|nous allons|a faire|action|relancer|envoyer|preparer|valider)\b"
    )
    decision_pattern = re_compile(r"\b(decision|decide|valide|accord|on part sur|retenu|approuve)\b")
    question_pattern = re_compile(r"\?")

    notes = {"actions": [], "decisions": [], "questions": []}
    for turn in turns:
        text = turn["text"]
        item = {
            "speaker": turn["speaker"],
            "time": format_ts(turn["start"], sep="."),
            "text": text,
        }
        lowered = strip_accents(text.lower())
        if action_pattern.search(lowered):
            notes["actions"].append(item)
        if decision_pattern.search(lowered):
            notes["decisions"].append(item)
        if question_pattern.search(text):
            notes["questions"].append(item)
    return notes


def strip_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    return "".join(character for character in normalized if not unicodedata.combining(character))


def re_compile(pattern: str) -> Any:
    return re.compile(pattern, re.IGNORECASE)


def write_notes_markdown(path: Path, turns: list[dict[str, Any]], source: Path) -> None:
    notes = extract_meeting_notes(turns)
    speakers = sorted({turn["speaker"] for turn in turns})
    duration = format_ts(turns[-1]["end"], sep=".") if turns else "00:00:00.000"

    with path.open("w", encoding="utf-8") as handle:
        handle.write("# Notes de transcription\n\n")
        handle.write(f"Source: `{source.name}`\n\n")
        handle.write(f"- Duree: `{duration}`\n")
        handle.write(f"- Locuteurs detectes: {', '.join(speakers) if speakers else 'aucun'}\n")
        handle.write(f"- Tours de parole: {len(turns)}\n\n")

        sections = (
            ("Actions possibles", notes["actions"]),
            ("Decisions possibles", notes["decisions"]),
            ("Questions", notes["questions"]),
        )
        for title, items in sections:
            handle.write(f"## {title}\n\n")
            if not items:
                handle.write("_Aucun element detecte automatiquement._\n\n")
                continue
            for item in items:
                handle.write(f'- `{item["time"]}` **{item["speaker"]}**: {item["text"]}\n')
            handle.write("\n")


def write_docx(path: Path, turns: list[dict[str, Any]], source: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        print("DOCX export skipped: install python-docx to enable it.")
        return False

    document = Document()
    document.add_heading("Transcription", level=1)
    document.add_paragraph(f"Source: {source.name}")
    for turn in turns:
        paragraph = document.add_paragraph()
        speaker = paragraph.add_run(f'{turn["speaker"]} ')
        speaker.bold = True
        paragraph.add_run(f'[{format_ts(turn["start"], sep=".")} - {format_ts(turn["end"], sep=".")}]')
        document.add_paragraph(turn["text"])
    document.save(path)
    return True


def parse_speaker_map(value: str) -> dict[str, str]:
    mapping: dict[str, str] = {}
    if not value.strip():
        return mapping
    for item in value.split(","):
        if not item.strip():
            continue
        if "=" not in item:
            fail(f"Invalid speaker mapping '{item}'. Expected SPEAKER_00=Name.")
        key, name = item.split("=", 1)
        key = key.strip()
        name = name.strip()
        if not key or not name:
            fail(f"Invalid speaker mapping '{item}'. Expected SPEAKER_00=Name.")
        mapping[key] = name
    return mapping


def load_speaker_map(args: argparse.Namespace) -> dict[str, str]:
    mapping = parse_speaker_map(args.speaker_map)
    if args.speaker_map_file:
        path = Path(args.speaker_map_file).expanduser()
        with path.open("r", encoding="utf-8") as handle:
            file_mapping = json.load(handle)
        if not isinstance(file_mapping, dict):
            fail("--speaker-map-file must contain a JSON object.")
        mapping.update({str(key): str(value) for key, value in file_mapping.items() if str(value).strip()})
    return mapping


def apply_speaker_map(result: dict[str, Any], mapping: dict[str, str]) -> dict[str, Any]:
    if not mapping:
        return result
    updated = json.loads(json.dumps(result))
    for segment in updated.get("segments", []):
        speaker = segment.get("speaker")
        if speaker in mapping:
            segment["speaker"] = mapping[speaker]
        for word in segment.get("words", []) or []:
            word_speaker = word.get("speaker")
            if word_speaker in mapping:
                word["speaker"] = mapping[word_speaker]
    updated["speaker_map"] = mapping
    return updated


def write_outputs(paths: Paths, result: dict[str, Any]) -> list[Path]:
    stem = transcript_stem(paths.audio)
    result = dict(result)
    result.pop("_checkpoint", None)
    result = sanitize_segments_for_alignment(result, "export")
    segments = normalize_segments(result)
    turns = merge_turns(segments)

    raw_json = paths.output_dir / f"{stem}.whisperx.json"
    segments_json = paths.output_dir / f"{stem}.segments.json"
    txt = paths.output_dir / f"{stem}.speaker-turns.txt"
    srt = paths.output_dir / f"{stem}.speaker-segments.srt"
    md = paths.output_dir / f"{stem}.speaker-turns.md"
    clean_txt = paths.output_dir / f"{stem}.clean.txt"
    notes_md = paths.output_dir / f"{stem}.notes.md"
    docx = paths.output_dir / f"{stem}.transcript.docx"

    with raw_json.open("w", encoding="utf-8") as handle:
        json.dump(result, handle, ensure_ascii=False, indent=2)
    with segments_json.open("w", encoding="utf-8") as handle:
        json.dump({"source": str(paths.audio), "segments": segments, "turns": turns}, handle, ensure_ascii=False, indent=2)
    write_txt(txt, turns)
    write_srt(srt, segments)
    write_markdown(md, turns, paths.audio)
    write_clean_txt(clean_txt, turns)
    write_notes_markdown(notes_md, turns, paths.audio)
    wrote_docx = write_docx(docx, turns, paths.audio)

    print("")
    print("Done. Files written:")
    output_paths = [txt, md, clean_txt, srt, segments_json, notes_md, raw_json]
    if wrote_docx:
        output_paths.insert(6, docx)
    for path in output_paths:
        print(f"- {path}")
    return output_paths


def append_history(paths: Paths, args: argparse.Namespace, output_paths: list[Path], status: str = "success") -> None:
    history = paths.output_dir / "transcription-history.jsonl"
    duration_seconds = None
    segments_path = paths.output_dir / f"{transcript_stem(paths.audio)}.segments.json"
    if segments_path.exists():
        try:
            with segments_path.open("r", encoding="utf-8") as handle:
                data = json.load(handle)
            turns = data.get("turns") or []
            if turns:
                duration_seconds = max(float(turn.get("end", 0.0)) for turn in turns)
        except Exception:
            duration_seconds = None
    record = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": status,
        "source_audio": str(paths.audio),
        "stem": transcript_stem(paths.audio),
        "duration_seconds": duration_seconds,
        "language": args.language,
        "profile": args.profile,
        "model": args.model,
        "asr_backend": args.asr_backend,
        "diarization": not args.no_diarization,
        "outputs": [str(path) for path in output_paths],
    }
    with history.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(record, ensure_ascii=False) + "\n")
    print(f"History updated: {history}")


def import_status(module: str) -> str:
    try:
        __import__(module)
        return "ok"
    except Exception as exc:
        return f"missing ({exc})"


def load_diarization_pipeline_class() -> Any:
    with warnings.catch_warnings():
        warnings.filterwarnings("ignore", message=".*torchcodec is not installed correctly.*")
        from whisperx.diarize import DiarizationPipeline

    return DiarizationPipeline


def run_doctor() -> None:
    print("Microwest Whisper doctor")
    print(f"- Python: {sys.version.split()[0]} ({sys.executable})")
    print(f"- Platform: {platform.system()} {platform.machine()}")
    print(f"- CPU cores: {os.cpu_count() or 'unknown'}")
    ffmpeg = find_ffmpeg()
    print(f"- FFmpeg: {ffmpeg or 'missing'}")
    print(f"- whisperx: {import_status('whisperx')}")
    print(f"- PySide6: {import_status('PySide6')}")
    print(f"- python-dotenv: {import_status('dotenv')}")
    print(f"- python-docx: {import_status('docx')}")
    print(f"- mlx-whisper: {import_status('mlx_whisper')}")
    try:
        import torch

        print(f"- torch: ok")
        print(f"- CUDA: {'available' if torch.cuda.is_available() else 'not available'}")
    except Exception as exc:
        print(f"- torch: missing ({exc})")


def check_hf_token(token: str | None) -> None:
    if not token:
        fail("Hugging Face token missing. Set HUGGINGFACE_TOKEN or pass --hf-token.")
    try:
        DiarizationPipeline = load_diarization_pipeline_class()
        DiarizationPipeline(token=token, device="cpu")
    except Exception as exc:
        fail(
            "Hugging Face token check failed. Verify the token and accept pyannote model terms. "
            f"Details: {exc}"
        )
    print("Hugging Face token check: ok")


def regenerate_from_checkpoint(paths: Paths, args: argparse.Namespace) -> list[Path]:
    stages = ["final.diarized", "final.no-diarization", "aligned"]
    if args.no_diarization:
        stages = ["final.no-diarization", "final.diarized", "aligned"]
    result = None
    for stage in stages:
        result = load_checkpoint(paths, stage, force=False)
        if result is not None:
            break
    if result is None:
        fail("No reusable checkpoint found. Run a transcription before --rename-only.")
    result = apply_speaker_map(result, load_speaker_map(args))
    return write_outputs(paths, result)


def transcribe(paths: Paths, args: argparse.Namespace) -> dict[str, Any]:
    import gc

    try:
        import whisperx
    except ImportError as exc:
        fail(f"WhisperX is not installed in this environment. Run the setup script first. Details: {exc}")

    backend = resolve_asr_backend(args.asr_backend)
    device = "cpu" if backend == "mlx" else select_device(args.device)

    final_checkpoint = "final.no-diarization" if args.no_diarization else "final.diarized"
    final_signature = checkpoint_signature(args, final_checkpoint, backend)
    asr_signature = checkpoint_signature(args, "asr", backend)
    aligned_signature = checkpoint_signature(args, "aligned", backend)

    final_result = load_checkpoint(paths, final_checkpoint, args.force, final_signature)
    if final_result is not None:
        return final_result

    audio = whisperx.load_audio(str(paths.wav))

    result = load_checkpoint(paths, "asr", args.force, asr_signature)
    if result is None:
        if backend == "mlx":
            result = transcribe_with_mlx(paths, args)
            device = "cpu"
        else:
            result, device = transcribe_with_whisperx(audio, args)
        result = sanitize_segments_for_alignment(result, "alignment")
        result["source_audio"] = str(paths.audio)
        result["preprocessed_wav"] = str(paths.wav)
        save_checkpoint(paths, "asr", result, asr_signature)

    result = sanitize_segments_for_alignment(result, "alignment")
    result["source_audio"] = str(paths.audio)
    result["preprocessed_wav"] = str(paths.wav)

    aligned = load_checkpoint(paths, "aligned", args.force, aligned_signature)
    if aligned is None:
        print("Aligning timestamps...")
        model_a, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
        aligned = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)
        aligned["source_audio"] = str(paths.audio)
        aligned["preprocessed_wav"] = str(paths.wav)
        aligned["asr_backend"] = result.get("asr_backend")
        save_checkpoint(paths, "aligned", aligned, aligned_signature)

        del model_a
        gc.collect()
    aligned = sanitize_segments_for_alignment(aligned, "export")
    result = aligned

    if args.no_diarization:
        print("Diarization skipped.")
        for segment in result.get("segments", []):
            segment["speaker"] = "SPEAKER_00"
        save_checkpoint(paths, final_checkpoint, result, final_signature)
        return result

    if not args.hf_token:
        fail("Hugging Face token missing. Pass --hf-token or set HUGGINGFACE_TOKEN.")

    print("Running speaker diarization with pyannote...")
    try:
        DiarizationPipeline = load_diarization_pipeline_class()
        diarize_model = DiarizationPipeline(token=args.hf_token, device=device)
    except Exception as exc:
        fail(
            "Speaker diarization could not start. Check that the Hugging Face token is valid and that "
            "the pyannote model terms have been accepted, especially "
            "https://huggingface.co/pyannote/speaker-diarization-community-1. Details: "
            f"{exc}"
        )
    diarize_kwargs: dict[str, int] = {}
    if args.speakers:
        diarize_kwargs["num_speakers"] = args.speakers
    if args.min_speakers:
        diarize_kwargs["min_speakers"] = args.min_speakers
    if args.max_speakers:
        diarize_kwargs["max_speakers"] = args.max_speakers

    try:
        diarize_segments = diarize_model(audio, **diarize_kwargs)
    except Exception as exc:
        fail(f"Speaker diarization failed while processing the audio. Details: {exc}")
    result = whisperx.assign_word_speakers(diarize_segments, result)
    result["source_audio"] = str(paths.audio)
    result["preprocessed_wav"] = str(paths.wav)
    save_checkpoint(paths, final_checkpoint, result, final_signature)
    return result


def main() -> None:
    args = parse_args()
    try:
        apply_profile(args)
        validate_args(args)
        if args.license_status:
            print(license_status_text())
            return
        if args.activate_license:
            result = activate_product_license(args.activate_license)
            if not result.ok:
                fail(result.message)
            print(result.message)
            print(license_status_text(result.state))
            return
        if args.doctor:
            run_doctor()
            return
        if args.check_token and not args.audio:
            check_hf_token(args.hf_token)
            return
        paths = build_paths(args)
        if args.check_token:
            check_hf_token(args.hf_token)
        license_check = validate_product_license()
        if not license_check.ok:
            fail(license_check.message)
        if args.rename_only:
            output_paths = regenerate_from_checkpoint(paths, args)
            append_history(paths, args, output_paths, status="renamed")
            return
        prepare_wav(paths, args)
        result = transcribe(paths, args)
        result = apply_speaker_map(result, load_speaker_map(args))
        output_paths = write_outputs(paths, result)
        append_history(paths, args, output_paths)
        if args.delete_work:
            paths.wav.unlink(missing_ok=True)
            print(f"Intermediate WAV deleted: {paths.wav}")
        else:
            print(f"Intermediate WAV kept for reuse: {paths.wav}")
    except subprocess.CalledProcessError as exc:
        command = " ".join(str(part) for part in exc.cmd) if exc.cmd else "external command"
        fail(f"Command failed with exit code {exc.returncode}: {command}")
    except KeyboardInterrupt:
        fail("Interrupted by user.")


if __name__ == "__main__":
    main()
